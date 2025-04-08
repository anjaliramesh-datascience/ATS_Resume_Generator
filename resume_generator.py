#!/usr/bin/env python3
"""
Resume Generator for ATS Optimization
This script creates ATS-friendly resumes optimized to score 90+ in Applicant Tracking Systems.
"""

import os
import re
import argparse
import json
from datetime import datetime
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import docx.opc.constants
import tempfile
import sys  # Added for sys.exit

# Add this function to create hyperlinks in Word documents
def add_hyperlink(paragraph, url, text, color=None, underline=True):
    """
    Add a hyperlink to a paragraph.
    
    :param paragraph: The paragraph to add the hyperlink to
    :param url: The URL to link to
    :param text: The text to display for the link
    :param color: The color of the link (RGB tuple), or None for default
    :param underline: Whether to underline the link
    :return: The hyperlink run
    """
    # This gets access to the document.xml.rels file and adds a new relationship ID
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run element
    new_run = OxmlElement('w:r')
    
    # Create new run properties element
    rPr = OxmlElement('w:rPr')
    
    # Add color if specified
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    
    # Add hyperlink styling (usually blue + underline)
    if not color:
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Default hyperlink color (blue)
        rPr.append(color)
    
    # Add underline if specified
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    # Join all the xml elements
    new_run.append(rPr)
    
    # Set the text of the hyperlink
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

class ResumeGenerator:
    """Generate ATS-optimized resumes"""
    
    def __init__(self, name, contact_info, professional_summary=None):
        """Initialize resume generator with personal details"""
        self.name = name
        self.contact_info = contact_info
        self.professional_summary = professional_summary
        self.work_experience = []
        self.technical_skills = {}
        self.education = []
        # Add optional sections
        self.internships = []
        self.projects = []
        self.certifications = []
        self.doc = Document()
        
        # Set font size constraints
        self.min_font_size = 12
        self.max_font_size = 14
        self.min_margin = 0.5  # in inches
        self.max_margin = 2.0  # in inches
        
        # Default font sizes - these will be dynamically adjusted if needed
        self.name_font_size = self.max_font_size + 2  # Name is slightly larger
        self.heading_font_size = self.max_font_size
        self.normal_font_size = self.min_font_size
        self.margin_size = 0.75  # in inches - start with reasonable margins
        
        self._setup_document()
    
    def _setup_document(self):
        """Set up document styles and formatting for ATS optimization"""
        # Set up margins (will be adjusted dynamically if needed)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(self.margin_size)
            section.bottom_margin = Inches(self.margin_size)
            section.left_margin = Inches(self.margin_size)
            section.right_margin = Inches(self.margin_size)
        
        # Create styles with current font sizes
        self._create_styles()
    
    def _create_styles(self):
        """Create document styles with current font sizes"""
        styles = self.doc.styles
        
        # Heading style
        heading_style = styles.add_style('ATS Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(self.heading_font_size)
        heading_style.font.bold = True
        heading_paragraph_format = heading_style.paragraph_format
        heading_paragraph_format.space_before = Pt(6)
        heading_paragraph_format.space_after = Pt(2)
        
        # Normal text style
        normal_style = styles.add_style('ATS Normal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(self.normal_font_size)
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.space_before = Pt(0)
        normal_paragraph_format.space_after = Pt(0)
        
        # Add compressed line spacing
        normal_paragraph_format.line_spacing = 1.0
        
        # Name style
        name_style = styles.add_style('Name Style', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(self.name_font_size)
        name_style.font.bold = True
        paragraph_format = name_style.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_after = Pt(2)
    
    def add_work_experience(self, title, company, location, start_date, end_date, responsibilities, achievements=None):
        """Add work experience entry"""
        self.work_experience.append({
            'title': title,
            'company': company,
            'location': location,
            'start_date': start_date,
            'end_date': end_date, 
            'responsibilities': responsibilities,
            'achievements': achievements or []
        })
    
    def add_technical_skills(self, category, skills):
        """Add technical skills by category"""
        self.technical_skills[category] = skills
    
    def add_education(self, degree, institution, location, graduation_date, gpa=None, relevant_courses=None):
        """Add education entry"""
        self.education.append({
            'degree': degree,
            'institution': institution,
            'location': location,
            'graduation_date': graduation_date,
            'gpa': gpa,
            'relevant_courses': relevant_courses
        })
    
    def add_internship(self, title, company, location, start_date, end_date, responsibilities, achievements=None):
        """Add internship entry"""
        self.internships.append({
            'title': title,
            'company': company,
            'location': location,
            'start_date': start_date,
            'end_date': end_date, 
            'responsibilities': responsibilities,
            'achievements': achievements or []
        })
    
    def add_project(self, name, description, technologies, url=None, start_date=None, end_date=None):
        """Add project entry"""
        self.projects.append({
            'name': name,
            'description': description,
            'technologies': technologies,
            'url': url,
            'start_date': start_date,
            'end_date': end_date
        })
    
    def add_certification(self, name, issuer, date, expiration_date=None, url=None):
        """Add certification entry"""
        self.certifications.append({
            'name': name,
            'issuer': issuer,
            'date': date,
            'expiration_date': expiration_date,
            'url': url
        })
    
    def _add_header(self):
        """Add name and contact information at the top of the resume"""
        # Add name
        name_paragraph = self.doc.add_paragraph(self.name, style='Name Style')
        
        # Add contact information - separating LinkedIn to make it clickable
        contact_paragraph = self.doc.add_paragraph(style='ATS Normal')
        contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process contact info to identify and make LinkedIn link clickable
        linkedin_url = None
        other_contacts = []
        
        for item in self.contact_info:
            if 'linkedin.com' in item.lower():
                linkedin_url = item
            else:
                other_contacts.append(item)
        
        # Add non-LinkedIn contact info
        if other_contacts:
            contact_text = " | ".join(other_contacts)
            if linkedin_url:  # Add a separator if LinkedIn will follow
                contact_text += " | "
            contact_paragraph.add_run(contact_text)
        
        # Add LinkedIn as hyperlink if it exists
        if linkedin_url:
            # Clean the LinkedIn URL if needed to ensure it has proper format
            if not linkedin_url.startswith('http'):
                url = f"https://{linkedin_url}"
            else:
                url = linkedin_url
                
            # Add the hyperlink using our custom function
            add_hyperlink(contact_paragraph, url, linkedin_url, '0000FF', True)
        
        # Remove extra paragraph space
        contact_paragraph.paragraph_format.space_after = Pt(4)
    
    def _add_professional_summary(self):
        """Add professional summary section"""
        if self.professional_summary:
            heading = self.doc.add_paragraph("PROFESSIONAL SUMMARY", style='ATS Heading')
            
            summary = self.doc.add_paragraph(self.professional_summary, style='ATS Normal')
            summary.paragraph_format.space_after = Pt(4)
    
    def _add_work_experience(self):
        """Add work experience section"""
        heading = self.doc.add_paragraph("WORK EXPERIENCE", style='ATS Heading')
        
        for i, job in enumerate(self.work_experience):
            # Job title and company
            job_heading = self.doc.add_paragraph(style='ATS Normal')
            job_title = job_heading.add_run(f"{job['title']} | {job['company']}")
            job_title.bold = True
            
            # Location and dates - same line to save space
            job_info = job_heading
            job_info.add_run(f" — {job['location']} | {job['start_date']} - {job['end_date']}")
            job_info.paragraph_format.space_after = Pt(2)
            
            # Responsibilities
            for resp in job['responsibilities']:
                resp_para = self.doc.add_paragraph(style='ATS Normal')
                resp_para.paragraph_format.left_indent = Inches(0.2)
                resp_para.paragraph_format.space_after = Pt(0)
                resp_para.add_run("• " + resp)
            
            # Achievements
            if job['achievements']:
                for achievement in job['achievements']:
                    ach_para = self.doc.add_paragraph(style='ATS Normal')
                    ach_para.paragraph_format.left_indent = Inches(0.2)
                    ach_para.paragraph_format.space_after = Pt(0)
                    ach_para.add_run("✓ " + achievement)
            
            # Only add space between jobs, not after the last one
            if i < len(self.work_experience) - 1:
                self.doc.add_paragraph(style='ATS Normal').paragraph_format.space_after = Pt(2)
    
    def _add_internships(self):
        """Add internships section if any exist"""
        if not self.internships:
            return  # Skip section if no internships
            
        heading = self.doc.add_paragraph("INTERNSHIPS", style='ATS Heading')
        
        for i, internship in enumerate(self.internships):
            # Internship title and company
            intern_heading = self.doc.add_paragraph(style='ATS Normal')
            intern_title = intern_heading.add_run(f"{internship['title']} | {internship['company']}")
            intern_title.bold = True
            
            # Location and dates - same line to save space
            intern_info = intern_heading
            intern_info.add_run(f" — {internship['location']} | {internship['start_date']} - {internship['end_date']}")
            intern_info.paragraph_format.space_after = Pt(2)
            
            # Responsibilities
            for resp in internship['responsibilities']:
                resp_para = self.doc.add_paragraph(style='ATS Normal')
                resp_para.paragraph_format.left_indent = Inches(0.2)
                resp_para.paragraph_format.space_after = Pt(0)
                resp_para.add_run("• " + resp)
            
            # Achievements
            if internship['achievements']:
                for achievement in internship['achievements']:
                    ach_para = self.doc.add_paragraph(style='ATS Normal')
                    ach_para.paragraph_format.left_indent = Inches(0.2)
                    ach_para.paragraph_format.space_after = Pt(0)
                    ach_para.add_run("✓ " + achievement)
            
            # Only add space between internships, not after the last one
            if i < len(self.internships) - 1:
                self.doc.add_paragraph(style='ATS Normal').paragraph_format.space_after = Pt(2)
    
    def _add_projects(self):
        """Add projects section if any exist"""
        if not self.projects:
            return  # Skip section if no projects
            
        heading = self.doc.add_paragraph("PROJECTS", style='ATS Heading')
        
        for i, project in enumerate(self.projects):
            # Project name with optional URL as hyperlink
            proj_heading = self.doc.add_paragraph(style='ATS Normal')
            proj_name = proj_heading.add_run(f"{project['name']}")
            proj_name.bold = True
            
            # Add date range if provided
            if project['start_date'] and project['end_date']:
                proj_heading.add_run(f" | {project['start_date']} - {project['end_date']}")
            
            # Add URL as hyperlink if provided
            if project['url']:
                proj_heading.add_run(" | ")
                url = project['url']
                if not url.startswith('http'):
                    url = f"https://{url}"
                add_hyperlink(proj_heading, url, "Project Link", '0000FF', True)
            
            # Project description
            desc_para = self.doc.add_paragraph(style='ATS Normal')
            desc_para.paragraph_format.left_indent = Inches(0.2)
            desc_para.add_run(project['description'])
            
            # Technologies used
            if project['technologies']:
                tech_para = self.doc.add_paragraph(style='ATS Normal')
                tech_para.paragraph_format.left_indent = Inches(0.2)
                tech_para.paragraph_format.space_after = Pt(0)
                tech_text = tech_para.add_run("Technologies: ")
                tech_text.bold = True
                tech_para.add_run(", ".join(project['technologies']))
            
            # Add space between projects, not after the last one
            if i < len(self.projects) - 1:
                self.doc.add_paragraph(style='ATS Normal').paragraph_format.space_after = Pt(2)
    
    def _add_certifications(self):
        """Add certifications section if any exist"""
        if not self.certifications:
            return  # Skip section if no certifications
            
        heading = self.doc.add_paragraph("CERTIFICATIONS", style='ATS Heading')
        
        for i, cert in enumerate(self.certifications):
            # Certification name and issuer
            cert_heading = self.doc.add_paragraph(style='ATS Normal')
            cert_name = cert_heading.add_run(f"{cert['name']} | {cert['issuer']}")
            cert_name.bold = True
            
            # Date and expiration if available
            date_info = f" | Issued: {cert['date']}"
            if cert['expiration_date']:
                date_info += f" | Expires: {cert['expiration_date']}"
            cert_heading.add_run(date_info)
            
            # Add URL as hyperlink if provided
            if cert['url']:
                cert_heading.add_run(" | ")
                url = cert['url']
                if not url.startswith('http'):
                    url = f"https://{url}"
                add_hyperlink(cert_heading, url, "Verify", '0000FF', True)
            
            # Add space between certifications, not after the last one
            if i < len(self.certifications) - 1:
                cert_heading.paragraph_format.space_after = Pt(2)
    
    def _add_technical_skills(self):
        """Add technical skills section in a more compact format"""
        heading = self.doc.add_paragraph("TECHNICAL SKILLS", style='ATS Heading')
        
        # Use a table-like format to fit more skills in less space
        skills_para = self.doc.add_paragraph(style='ATS Normal')
        
        # Process all skills into a single paragraph with proper formatting
        for i, (category, skills) in enumerate(self.technical_skills.items()):
            category_text = skills_para.add_run(f"{category}: ")
            category_text.bold = True
            skills_para.add_run(", ".join(skills))
            
            # Add line break between categories except for the last one
            if i < len(self.technical_skills) - 1:
                skills_para.add_run("\n")
        
        skills_para.paragraph_format.space_after = Pt(4)
    
    def _add_education(self):
        """Add education section in a more compact format"""
        heading = self.doc.add_paragraph("EDUCATION", style='ATS Heading')
        
        for i, edu in enumerate(self.education):
            # Degree and institution
            edu_heading = self.doc.add_paragraph(style='ATS Normal')
            degree_text = edu_heading.add_run(f"{edu['degree']} | {edu['institution']}")
            degree_text.bold = True
            
            # Location and graduation date - same paragraph
            edu_heading.add_run(f" — {edu['location']} | {edu['graduation_date']}")
            
            # Add GPA if available
            if edu['gpa']:
                edu_heading.add_run(f" (GPA: {edu['gpa']})")
            
            # Add relevant courses if available
            if edu['relevant_courses'] and len(edu['relevant_courses']) > 0:
                courses_para = self.doc.add_paragraph(style='ATS Normal')
                courses_para.paragraph_format.left_indent = Inches(0.2)
                courses_text = courses_para.add_run("Relevant Coursework: ")
                courses_text.bold = True
                courses_para.add_run(", ".join(edu['relevant_courses']))
            
            # Only add space between education entries, not after the last one
            if i < len(self.education) - 1:
                edu_heading.paragraph_format.space_after = Pt(2)
    
    def _estimate_content_volume(self):
        """
        Estimate the total amount of content in the resume
        Returns a score representing approximate text volume
        """
        volume = 0
        
        # Count professional summary
        if self.professional_summary:
            volume += len(self.professional_summary) * 0.5
        
        # Count work experience (most significant content)
        for job in self.work_experience:
            # Job title, company, location info
            volume += 100  # Base value for each job
            
            # Add for responsibilities
            for resp in job['responsibilities']:
                volume += len(resp) * 0.7
            
            # Add for achievements
            for achievement in job['achievements']:
                volume += len(achievement) * 0.7
        
        # Count skills (less space-intensive)
        for category, skills in self.technical_skills.items():
            volume += len(category) + sum(len(skill) for skill in skills) * 0.3
        
        # Count education
        for edu in self.education:
            volume += 80  # Base value for each education entry
        
        return volume
    
    def _adjust_font_and_margins(self):
        """
        Dynamically adjust font sizes and margins based on content volume
        while respecting minimum/maximum constraints
        """
        content_volume = self._estimate_content_volume()
        
        # Define adjustment thresholds and scaling factors
        # These values are calibrated based on typical resume content lengths
        if content_volume > 1800:  # Very large resume
            # Use minimum font sizes and margins to fit more content
            self.name_font_size = self.max_font_size
            self.heading_font_size = self.min_font_size + 1  # 13pt
            self.normal_font_size = self.min_font_size  # 12pt
            self.margin_size = self.min_margin  # 0.5 inches
        elif content_volume > 1500:  # Large resume
            self.name_font_size = self.max_font_size + 1  # 15pt
            self.heading_font_size = self.min_font_size + 1  # 13pt
            self.normal_font_size = self.min_font_size  # 12pt
            self.margin_size = 0.6  # 0.6 inches
        elif content_volume > 1200:  # Medium-large resume
            self.name_font_size = self.max_font_size + 2  # 16pt
            self.heading_font_size = self.max_font_size - 1  # 13pt
            self.normal_font_size = self.min_font_size  # 12pt
            self.margin_size = 0.7  # 0.7 inches
        elif content_volume > 900:  # Medium resume
            self.name_font_size = self.max_font_size + 2  # 16pt
            self.heading_font_size = self.max_font_size  # 14pt
            self.normal_font_size = self.min_font_size  # 12pt
            self.margin_size = 0.9  # 0.9 inches
        else:  # Small resume - use maximum sizes for better readability
            self.name_font_size = self.max_font_size + 2  # 16pt
            self.heading_font_size = self.max_font_size  # 14pt
            self.normal_font_size = self.min_font_size  # 12pt
            self.margin_size = 1.2  # Larger margins for smaller resumes
            
        print(f"Content volume: {content_volume}")
        print(f"Adjusted fonts - Name: {self.name_font_size}pt, Heading: {self.heading_font_size}pt, Normal: {self.normal_font_size}pt")
        print(f"Adjusted margins: {self.margin_size} inches")
        
        # Recreate the document with new sizes
        self.doc = Document()  # Create a fresh document
        self._setup_document()  # Apply the adjusted styles
    
    def _check_and_adjust_for_page_fit(self, output_filename):
        """
        Check if resume fits on one page, and adjust if needed
        Uses trial-and-error approach by saving temp files and checking
        """
        # First try with current settings
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_file.close()
        
        # Save current document to temp file
        self.doc.save(temp_file.name)
        
        # Check if it fits on one page by loading it back
        doc_check = Document(temp_file.name)
        
        # Count pages (this is an approximation since python-docx doesn't have direct page counting)
        page_count = self._estimate_page_count(doc_check)
        
        attempts = 0
        max_attempts = 3
        
        while page_count > 1 and attempts < max_attempts:
            attempts += 1
            
            # Reduce sizes by ~5% each attempt, respecting minimum constraints
            self.name_font_size = max(self.max_font_size, self.name_font_size * 0.95)
            self.heading_font_size = max(self.min_font_size + 1, self.heading_font_size * 0.95)
            self.normal_font_size = self.min_font_size  # Don't reduce below minimum
            
            # Reduce margins, respecting minimum margin
            self.margin_size = max(self.min_margin, self.margin_size * 0.9)
            
            # Recreate document with adjusted sizes
            self.doc = Document()
            self._setup_document()
            
            # Regenerate content
            self._add_header()
            self._add_professional_summary()
            self._add_work_experience()
            self._add_internships()
            self._add_projects()
            self._add_certifications()
            self._add_technical_skills()
            self._add_education()
            
            # Save to temp and check again
            self.doc.save(temp_file.name)
            doc_check = Document(temp_file.name)
            page_count = self._estimate_page_count(doc_check)
            
            print(f"Adjustment attempt {attempts}: estimated {page_count} pages")
            print(f"  • Adjusted to - Name: {self.name_font_size}pt, Heading: {self.heading_font_size}pt, Normal: {self.normal_font_size}pt")
            print(f"  • Margins: {self.margin_size} inches")
        
        # If still not fitting after max attempts, warn the user
        if page_count > 1:
            print("WARNING: Resume may extend to multiple pages. Consider reducing content or further adjusting settings.")
        
        # Clean up temp file
        os.unlink(temp_file.name)
        
        # Save final version to the actual output file
        self.doc.save(output_filename)
    
    def _estimate_page_count(self, doc):
        """
        Estimate number of pages in the document
        This is a heuristic since python-docx doesn't have direct page counting
        """
        # Get page size in points (1/72 of an inch)
        # Assuming US Letter size: 8.5 x 11 inches
        section = doc.sections[0]
        page_height_inches = 11  # Standard US Letter
        page_width_inches = 8.5
        
        # Convert to points
        page_height_pt = page_height_inches * 72
        
        # Account for margins
        margin_top = section.top_margin.pt
        margin_bottom = section.bottom_margin.pt
        usable_height_pt = page_height_pt - margin_top - margin_bottom
        
        # Estimate total height needed for all content
        total_content_height = 0
        
        # Estimate based on paragraph count and font sizes
        for paragraph in doc.paragraphs:
            # Base height for a paragraph
            if 'Heading' in paragraph.style.name:
                para_height = self.heading_font_size * 1.5  # Slightly more for headings
            elif 'Name' in paragraph.style.name:
                para_height = self.name_font_size * 1.5
            else:
                para_height = self.normal_font_size * 1.2
            
            # Add estimated line wrapping
            text_length = len(paragraph.text)
            chars_per_line = int((page_width_inches - section.left_margin.inches - section.right_margin.inches) * 12)  # ~12 chars per inch
            
            # Simple line wrapping estimate
            if chars_per_line > 0:
                lines = max(1, text_length // chars_per_line)
                para_height *= lines
            
            total_content_height += para_height
        
        # Estimate page count
        estimated_pages = max(1, total_content_height / usable_height_pt)
        return estimated_pages

    def generate_resume(self, output_filename=None):
        """Generate the resume document with dynamic sizing"""
        if not output_filename:
            output_filename = f"{self.name.replace(' ', '_')}_Resume.docx"
        
        # Adjust font sizes and margins based on content volume
        self._adjust_font_and_margins()
        
        # Add sections
        self._add_header()
        self._add_professional_summary()
        self._add_work_experience()
        self._add_internships()
        self._add_projects()
        self._add_certifications()
        self._add_technical_skills()
        self._add_education()
        
        # Check if resume fits on one page and adjust if needed
        self._check_and_adjust_for_page_fit(output_filename)
        
        print(f"Resume successfully generated: {output_filename}")
        return output_filename
    
    @staticmethod
    def ats_keyword_analysis(job_description, resume_text):
        """
        Analyze how well the resume matches a job description
        Returns a score and missing keywords
        """
        # Extract keywords from job description (simplified version)
        # In a real scenario, this would use more sophisticated NLP techniques
        job_words = re.findall(r'\b[a-zA-Z][a-zA-Z]+\b', job_description.lower())
        job_keywords = [word for word in set(job_words) if len(word) > 3]
        
        # Count matches
        matched_keywords = []
        missing_keywords = []
        
        for keyword in job_keywords:
            if keyword in resume_text.lower():
                matched_keywords.append(keyword)
            else:
                missing_keywords.append(keyword)
        
        # Calculate score (simplified)
        if len(job_keywords) > 0:
            match_percentage = (len(matched_keywords) / len(job_keywords)) * 100
        else:
            match_percentage = 0
            
        return {
            'score': match_percentage,
            'matched_keywords': matched_keywords,
            'missing_keywords': missing_keywords
        }


def load_example_data():
    """Load example data for demonstration"""
    return {
        "name": "John Doe",
        "contact_info": [
            "john.doe@email.com",
            "(555) 123-4567",
            "linkedin.com/in/johndoe",
            "San Francisco, CA"
        ],
        "professional_summary": "Results-driven Software Engineer with 5+ years of experience designing and developing scalable applications. Proficient in Python, JavaScript, and cloud technologies. Strong problem-solving skills and passion for creating efficient, maintainable code.",
        "work_experience": [
            {
                "title": "Senior Software Engineer",
                "company": "Tech Solutions Inc.",
                "location": "San Francisco, CA",
                "start_date": "January 2020",
                "end_date": "Present",
                "responsibilities": [
                    "Develop and maintain cloud-based applications using Python, Django, and AWS services",
                    "Lead a team of 5 engineers, implementing Agile methodologies and CI/CD practices",
                    "Optimize database queries and application performance, reducing load times by 40%"
                ],
                "achievements": [
                    "Implemented microservices architecture that improved system scalability by 200%",
                    "Reduced infrastructure costs by 30% through AWS optimization"
                ]
            },
            {
                "title": "Software Engineer",
                "company": "WebDev Enterprises",
                "location": "Oakland, CA",
                "start_date": "June 2018",
                "end_date": "December 2019",
                "responsibilities": [
                    "Developed responsive web applications using React, Node.js, and MongoDB",
                    "Collaborated with product managers to define requirements and features",
                    "Implemented automated testing, achieving 90% code coverage"
                ],
                "achievements": [
                    "Developed a feature that increased user engagement by 25%",
                    "Mentored 3 junior developers who were later promoted"
                ]
            }
        ],
        "technical_skills": {
            "Programming Languages": ["Python", "JavaScript", "TypeScript", "SQL", "HTML", "CSS"],
            "Frameworks & Libraries": ["Django", "Flask", "React", "Node.js", "Express", "Redux"],
            "Cloud & DevOps": ["AWS", "Docker", "Kubernetes", "CI/CD", "Git", "GitHub Actions"],
            "Databases": ["PostgreSQL", "MongoDB", "Redis", "DynamoDB"],
            "Tools & Methodologies": ["Agile", "Scrum", "Jira", "RESTful APIs", "GraphQL"]
        },
        "education": [
            {
                "degree": "Master of Science in Computer Science",
                "institution": "University of California, Berkeley",
                "location": "Berkeley, CA",
                "graduation_date": "May 2018",
                "gpa": "3.9/4.0",
                "relevant_courses": ["Advanced Algorithms", "Machine Learning", "Distributed Systems", "Cloud Computing"]
            },
            {
                "degree": "Bachelor of Science in Computer Engineering",
                "institution": "Stanford University",
                "location": "Stanford, CA",
                "graduation_date": "May 2016",
                "gpa": "3.8/4.0",
                "relevant_courses": ["Data Structures", "Computer Architecture", "Operating Systems", "Database Systems"]
            }
        ]
    }


def load_resume_from_json(json_file):
    """Load resume data from a JSON file"""
    try:
        with open(json_file, 'r') as file:
            data = json.load(file)
        
        # Validate required fields
        required_fields = ['name', 'contact_info', 'work_experience', 'technical_skills', 'education']
        for field in required_fields:
            if field not in data:
                raise ValueError(f"Missing required field in JSON file: {field}")
        
        print(f"Successfully loaded resume data from {json_file}")
        return data
    except json.JSONDecodeError:
        print(f"Error: {json_file} is not a valid JSON file")
        sys.exit(1)
    except Exception as e:
        print(f"Error loading resume data: {str(e)}")
        sys.exit(1)


def interactive_resume_builder():
    """Interactive CLI for building resume"""
    print("=== ATS-Optimized Resume Generator ===")
    
    # Personal Information
    name = input("Enter your full name: ")
    email = input("Enter your email: ")
    phone = input("Enter your phone number: ")
    linkedin = input("Enter your LinkedIn URL: ")
    location = input("Enter your location (City, State): ")
    contact_info = [email, phone, linkedin, location]
    
    summary = input("Enter your professional summary (press Enter to skip): ")
    
    # Initialize resume
    resume = ResumeGenerator(name, contact_info, summary if summary else None)
    
    # Work Experience
    print("\n=== Work Experience ===")
    more_experience = True
    while more_experience:
        title = input("Job title: ")
        company = input("Company name: ")
        location = input("Job location: ")
        start_date = input("Start date (e.g., January 2020): ")
        end_date = input("End date (or 'Present'): ")
        
        responsibilities = []
        print("Enter job responsibilities (one per line, blank line to finish):")
        while True:
            resp = input("- ")
            if not resp:
                break
            responsibilities.append(resp)
        
        achievements = []
        print("Enter key achievements (one per line, blank line to finish):")
        while True:
            ach = input("- ")
            if not ach:
                break
            achievements.append(ach)
        
        resume.add_work_experience(title, company, location, start_date, end_date, responsibilities, achievements)
        
        more = input("Add another job? (y/n): ").lower()
        more_experience = more.startswith('y')
    
    # Technical Skills
    print("\n=== Technical Skills ===")
    more_skills = True
    while more_skills:
        category = input("Skill category (e.g., Programming Languages): ")
        print(f"Enter {category} (comma-separated list):")
        skills_input = input("> ")
        skills = [s.strip() for s in skills_input.split(',')]
        
        resume.add_technical_skills(category, skills)
        
        more = input("Add another skill category? (y/n): ").lower()
        more_skills = more.startswith('y')
    
    # Education
    print("\n=== Education ===")
    more_education = True
    while more_education:
        degree = input("Degree/certification: ")
        institution = input("Institution: ")
        location = input("Location: ")
        graduation = input("Graduation date: ")
        gpa = input("GPA (optional): ")
        
        # Get relevant courses
        courses = []
        print("Enter relevant courses (comma-separated list):")
        courses_input = input("> ")
        if courses_input:
            courses = [course.strip() for course in courses_input.split(',')]
        
        resume.add_education(
            degree, 
            institution, 
            location, 
            graduation,
            gpa if gpa else None,
            courses if courses else None
        )
        
        more = input("Add another education entry? (y/n): ").lower()
        more_education = more.startswith('y')
    
    # Generate resume
    output_file = input("\nOutput filename (default is YourName_Resume.docx): ")
    if not output_file:
        output_file = None
    
    resume.generate_resume(output_file)


def create_resume_from_json(json_file):
    """Create a resume from JSON data"""
    data = load_resume_from_json(json_file)
    
    # Create resume object
    resume = ResumeGenerator(data['name'], data['contact_info'], 
                           data.get('professional_summary', None))
    
    # Add work experience
    for job in data['work_experience']:
        resume.add_work_experience(
            job['title'], 
            job['company'], 
            job['location'], 
            job['start_date'], 
            job['end_date'], 
            job['responsibilities'], 
            job.get('achievements', [])  # Some jobs might not have achievements
        )
    
    # Add internships if present
    if 'internships' in data:
        for internship in data['internships']:
            resume.add_internship(
                internship['title'],
                internship['company'],
                internship['location'],
                internship['start_date'],
                internship['end_date'],
                internship['responsibilities'],
                internship.get('achievements', [])
            )
    
    # Add projects if present
    if 'projects' in data:
        for project in data['projects']:
            resume.add_project(
                project['name'],
                project['description'],
                project['technologies'],
                project.get('url', None),
                project.get('start_date', None),
                project.get('end_date', None)
            )
    
    # Add technical skills
    for category, skills in data['technical_skills'].items():
        resume.add_technical_skills(category, skills)
    
    # Add certifications if present
    if 'certifications' in data:
        for cert in data['certifications']:
            resume.add_certification(
                cert['name'],
                cert['issuer'],
                cert['date'],
                cert.get('expiration_date', None),
                cert.get('url', None)
            )
    
    # Add education
    for edu in data['education']:
        resume.add_education(
            edu['degree'],
            edu['institution'],
            edu['location'],
            edu['graduation_date'],
            edu.get('gpa', None),
            edu.get('relevant_courses', None)
        )
    
    # Generate the resume with optional filename from JSON
    output_file = data.get('output_filename', None)
    return resume.generate_resume(output_file)


def main():
    parser = argparse.ArgumentParser(description='Generate ATS-optimized resumes')
    parser.add_argument('--example', action='store_true', help='Generate an example resume')
    parser.add_argument('--interactive', action='store_true', help='Use interactive mode to build resume')
    parser.add_argument('--json', metavar='FILE', help='Path to JSON file with resume data')
    
    args = parser.parse_args()
    
    if args.example:
        print("Generating example resume...")
        data = load_example_data()
        
        resume = ResumeGenerator(data['name'], data['contact_info'], data['professional_summary'])
        
        for job in data['work_experience']:
            resume.add_work_experience(
                job['title'], 
                job['company'], 
                job['location'], 
                job['start_date'], 
                job['end_date'], 
                job['responsibilities'], 
                job['achievements']
            )
        
        for category, skills in data['technical_skills'].items():
            resume.add_technical_skills(category, skills)
        
        for edu in data['education']:
            resume.add_education(
                edu['degree'],
                edu['institution'],
                edu['location'],
                edu['graduation_date'],
                edu.get('gpa', None),
                edu.get('relevant_courses', None)
            )
        
        resume.generate_resume()
    elif args.json:
        print(f"Creating resume from JSON file: {args.json}")
        create_resume_from_json(args.json)
    elif args.interactive:
        interactive_resume_builder()
    else:
        print("No mode selected. Use one of the following:")
        print("  --example      Generate an example resume")
        print("  --interactive  Use interactive mode to build resume")
        print("  --json FILE    Create resume from JSON data file")
        print("Run 'python resume-generator.py --help' for more information.")


if __name__ == "__main__":
    main()
